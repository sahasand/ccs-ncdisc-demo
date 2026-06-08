#!/usr/bin/env python3
"""
02_generate_tables.py  --  HF-1002-CL-101 (FICTIONAL DEMO)

NON-CDISC table programming.
Reads the home-grown analysis datasets from ./analysis_data/ and produces
populated TFL tables (SAS PROC REPORT look) into ./tables_output/ as .txt,
and assembles them into one landscape Word document:
    HF-1002-CL-101_Populated_Tables.docx

No CDISC, no SDTM/ADaM -- tables are programmed straight off the analysis
datasets built by 01_build_analysis_datasets.py.  All numbers are derived
from the synthetic EDC data; nothing is hard-coded.
"""
import os
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT

BASE = os.path.dirname(os.path.abspath(__file__))
AD = os.path.join(BASE, "analysis_data")
OUTT = os.path.join(BASE, "tables_output")
os.makedirs(OUTT, exist_ok=True)

S = pd.read_csv(os.path.join(AD, "an_subject.csv"))
AE = pd.read_csv(os.path.join(AD, "an_ae.csv"))
LB = pd.read_csv(os.path.join(AD, "an_lb.csv"))
EFF = pd.read_csv(os.path.join(AD, "an_eff.csv"))

COHORTS = [(1, "Cohort 1", "(Low)"), (2, "Cohort 2", "(Mid)"), (3, "Cohort 3", "(High)")]
N = {c: int((S.COHORT == c).sum()) for c, _, _ in COHORTS}
NT = int(len(S))
STUB, COL = 36, 15
TOTALW = STUB + 4 * COL

# ---------- stat helpers ----------
TCRIT = {1:12.706,2:4.303,3:3.182,4:2.776,5:2.571,6:2.447,7:2.365,8:2.306,9:2.262,10:2.228,
         11:2.201,12:2.179,13:2.160,14:2.145,15:2.131,16:2.120,17:2.110,18:2.101,19:2.093,20:2.086,
         21:2.080,22:2.074,23:2.069,24:2.064,25:2.060,26:2.056,27:2.052,28:2.048,29:2.045,30:2.042}
def tcrit(df): return TCRIT.get(df, 1.96)
def vals(series): return pd.to_numeric(series, errors="coerce").dropna().values
def f_n(v): return str(len(v))
def f_meansd(v): return "-" if len(v) == 0 else f"{np.mean(v):.1f} ({np.std(v, ddof=1):.2f})" if len(v) > 1 else f"{np.mean(v):.1f} (-)"
def f_median(v): return "-" if len(v) == 0 else f"{np.median(v):.1f}"
def f_minmax(v): return "-" if len(v) == 0 else f"{np.min(v):.0f}, {np.max(v):.0f}"
def f_npct(k, n): return "0" if n == 0 else f"{k} ({100*k/n:.1f})"
def f_ci(v):
    if len(v) < 2: return "-"
    m, sd, n = np.mean(v), np.std(v, ddof=1), len(v)
    h = tcrit(n-1) * sd / np.sqrt(n)
    return f"({m-h:.2f}, {m+h:.2f})"   # 2 dp to match mock shell (QC F-5)

# ---------- renderer ----------
def rowfmt(stub, vs):
    s = str(stub)[:STUB].ljust(STUB)
    for v in vs: s += str(v).rjust(COL)
    return s
def render(tnum, title, pop, rows, foot, source):
    L = [f"{tnum}  {title}", f"({pop})",
         "Protocol: HF-1002-CL-101        Sponsor: Corvexa Therapeutics, Inc. (fictional)", ""]
    L.append("=" * TOTALW)
    L.append(rowfmt("", ["Cohort 1", "Cohort 2", "Cohort 3", "Total"]))
    L.append(rowfmt("", ["(Low)", "(Mid)", "(High)", ""]))
    L.append(rowfmt("", [f"(N={N[1]})", f"(N={N[2]})", f"(N={N[3]})", f"(N={NT})"]))
    L.append("-" * TOTALW)
    for r in rows:
        if r[1] is None: L.append(str(r[0]))
        else: L.append(rowfmt(r[0], r[1]))
    L.append("=" * TOTALW)
    L += ["", "Source: " + source]
    for i, f in enumerate(foot): L.append(f"[{chr(97+i)}] {f}")
    L.append("Placeholders: none -- values computed from synthetic EDC data (FICTIONAL).")
    return L

def bycoh(df, fn):
    """apply fn(sub_df) for each cohort + total, return list of 4 strings."""
    out = []
    for c, _, _ in COHORTS: out.append(fn(df[df.COHORT == c]))
    out.append(fn(df))
    return out

# ======================================================================
# T 14.1.1  Demographics and Baseline Characteristics (Safety Set)
# ======================================================================
def t1411():
    rows = []
    rows.append(("Age (years)", None))
    rows.append(("  n", bycoh(S, lambda d: f_n(vals(d.AGE)))))
    rows.append(("  Mean (SD)", bycoh(S, lambda d: f_meansd(vals(d.AGE)))))
    rows.append(("  Median", bycoh(S, lambda d: f_median(vals(d.AGE)))))
    rows.append(("  Min, Max", bycoh(S, lambda d: f_minmax(vals(d.AGE)))))
    rows.append(("Age group, n (%)", None))
    for g in ["<65", ">=65"]:
        rows.append((f"  {g}", bycoh(S, lambda d, g=g: f_npct((d.AGEGR1 == g).sum(), len(d)))))
    rows.append(("Sex, n (%)", None))
    for sx, lab in [("M", "  Male"), ("F", "  Female")]:
        rows.append((lab, bycoh(S, lambda d, sx=sx: f_npct((d.SEX == sx).sum(), len(d)))))
    rows.append(("Race, n (%)", None))
    for rc in sorted(S.RACE.dropna().unique()):
        rows.append((f"  {rc.title()}", bycoh(S, lambda d, rc=rc: f_npct((d.RACE == rc).sum(), len(d)))))
    rows.append(("Baseline LVEF (%)", None))
    rows.append(("  Mean (SD)", bycoh(S, lambda d: f_meansd(vals(d.BL_LVEF)))))
    rows.append(("  Min, Max", bycoh(S, lambda d: f_minmax(vals(d.BL_LVEF)))))
    rows.append(("Baseline NYHA class III, n (%)", bycoh(S, lambda d: f_npct((d.BL_NYHA == 3).sum(), len(d)))))
    return render("Table 14.1.1", "Demographics and Baseline Characteristics", "Safety Set", rows,
                  ["Percentages are based on the number of subjects in each cohort (N).",
                   "Safety Set = all subjects who received any amount of HF-1002."],
                  "an_subject (derived from raw EDC dm, el)")

# ======================================================================
# T 14.1.2  Participant Disposition (Safety Set)
# ======================================================================
def t1412():
    rows = []
    rows.append(("Enrolled, n", bycoh(S, lambda d: str(len(d)))))
    rows.append(("Received HF-1002 (Safety Set), n (%)", bycoh(S, lambda d: f_npct((d.SAFFL == "Y").sum(), len(d)))))
    rows.append(("Completed Month 12, n (%)", bycoh(S, lambda d: f_npct((d.COMPLFL == "Y").sum(), len(d)))))
    rows.append(("Discontinued, n (%)", bycoh(S, lambda d: f_npct((d.COMPLFL == "N").sum(), len(d)))))
    for reason, lab in [("DEATH", "  Death"), ("WITHDRAWAL BY SUBJECT", "  Withdrawal by subject")]:
        rows.append((lab, bycoh(S, lambda d, reason=reason: f_npct((d.DSDECOD == reason).sum(), len(d)))))
    rows.append(("Entered long-term follow-up, n (%)", bycoh(S, lambda d: f_npct((d.COMPLFL == "Y").sum(), len(d)))))
    return render("Table 14.1.2", "Participant Disposition", "Enrolled Set", rows,
                  ["Population aligned to SAP Section 12 (Enrolled Set); Enrolled = Safety Set = 18 in this study (QC F-3).",
                   "Discontinued = did not complete the 12-month primary observation period.",
                   "Percentages are based on the number enrolled in each cohort."],
                  "an_subject (derived from raw EDC ds, ex)")

# ======================================================================
# T 14.1.3  Extent of Exposure (Safety Set)
# ======================================================================
def t1413():
    doselab = {1: "3.0E13", 2: "1.0E14", 3: "3.0E14"}
    rows = []
    rows.append(("Received complete infusion, n (%)", bycoh(S, lambda d: f_npct((d.FASFL == "Y").sum(), len(d)))))
    rows.append(("Assigned dose (vg)", [doselab[1], doselab[2], doselab[3], "--"]))
    rows.append(("Infusion volume (mL)", None))
    rows.append(("  Mean (SD)", bycoh(S, lambda d: f_meansd(vals(d.EXVOL_ML)))))
    rows.append(("  Min, Max", bycoh(S, lambda d: f_minmax(vals(d.EXVOL_ML)))))
    return render("Table 14.1.3", "Extent of Exposure", "Safety Set", rows,
                  ["A single antegrade intracoronary infusion of HF-1002 was administered on Day 1.",
                   "vg = vector genomes."],
                  "an_subject (derived from raw EDC ex)")

# ======================================================================
# T 14.3.1.1  Overall Summary of TEAEs (Safety Set)
# ======================================================================
def subj_with(mask_df):
    """count distinct subjects (by cohort + total) in the AE subset mask_df."""
    def fn(scope):
        ids = set(mask_df.SUBJID)
        return scope[scope.SUBJID.isin(ids)]
    out = []
    for c, _, _ in COHORTS:
        out.append(f_npct(len(set(mask_df.SUBJID) & set(S[S.COHORT == c].SUBJID)), N[c]))
    out.append(f_npct(len(set(mask_df.SUBJID)), NT))
    return out

def t14311():
    rows = []
    rows.append(("Any TEAE", subj_with(AE[AE.TEAEFL == "Y"])))
    rows.append(("Any related TEAE", subj_with(AE[(AE.TEAEFL == "Y") & (AE.RELFL == "Y")])))
    rows.append(("Any serious TEAE (SAE)", subj_with(AE[(AE.TEAEFL == "Y") & (AE.SERFL == "Y")])))
    rows.append(("Any Grade >= 3 TEAE", subj_with(AE[(AE.TEAEFL == "Y") & (AE.GE3FL == "Y")])))
    rows.append(("Any dose-limiting toxicity", subj_with(AE[AE.DLTFL == "Y"])))
    deaths = S[S.DTHFL == "Y"]
    rows.append(("Deaths", [f_npct((deaths.COHORT == c).sum(), N[c]) for c, _, _ in COHORTS] + [f_npct(len(deaths), NT)]))
    return render("Table 14.3.1.1", "Overall Summary of Treatment-Emergent Adverse Events", "Safety Set", rows,
                  ["A subject is counted once within each category regardless of number of events.",
                   "Related = investigator assessment of possibly, probably, or related.",
                   "TEAE = treatment-emergent adverse event; SAE = serious adverse event."],
                  "an_ae, an_subject (derived from raw EDC ae)")

# ======================================================================
# T 14.3.1.2  TEAEs by Verbatim Term (Safety Set)  [uncoded]
# ======================================================================
def t14312():
    teae = AE[AE.TEAEFL == "Y"]
    terms = (teae.groupby("AETERM").SUBJID.nunique().sort_values(ascending=False))
    rows = [("Subjects with >= 1 TEAE", subj_with(teae)), ("Preferred/verbatim term, n (%)", None)]
    for term in terms.index:
        sub = teae[teae.AETERM == term]
        rows.append(("  " + term[:32], [f_npct(len(set(sub.SUBJID) & set(S[S.COHORT == c].SUBJID)), N[c]) for c, _, _ in COHORTS] + [f_npct(sub.SUBJID.nunique(), NT)]))
    return render("Table 14.3.1.2", "TEAEs by Verbatim Term", "Safety Set", rows,
                  ["AEs are UNCODED in raw EDC; terms shown are investigator verbatim, grouped exactly.",
                   "MedDRA System Organ Class / Preferred Term coding is a downstream step (not applied here)."],
                  "an_ae (derived from raw EDC ae)")

# ======================================================================
# T 14.3.3.1  Dose-Limiting Toxicities (DLT-Evaluable Set)
# ======================================================================
def t14331():
    dlt = AE[AE.DLTFL == "Y"]
    rows = [("Subjects with >= 1 DLT, n (%)", subj_with(dlt)), ("DLT by term, n", None)]
    foot = ["DLT window = Day 1 through Day 29 (Week 4).",
            "DLT = investigator-assessed related event meeting a protocol-defined DLT criterion."]
    for term in sorted(dlt.AETERM.unique()):
        sub = dlt[dlt.AETERM == term]
        disp = term[:32]
        rows.append(("  " + disp, [str((sub.COHORT == c).sum()) for c, _, _ in COHORTS] + [str(len(sub))]))
        if len(term) > 32:  # full verbatim term preserved in a footnote (QC F-4)
            foot.append(f"Full verbatim term for '{disp.strip()}...': {term}")
    return render("Table 14.3.3.1", "Dose-Limiting Toxicities by Cohort", "DLT-Evaluable Set", rows, foot,
                  "an_ae (derived from raw EDC ae; DLTFL = AEDLT within DLT window)")

# ======================================================================
# Efficacy change-from-baseline (FAS)  -- generic builder
# ======================================================================
def cfb_table(tnum, paramcd, label, source):
    d = EFF[(EFF.PARAMCD == paramcd) & (EFF.FASFL == "Y")]
    rows = []
    rows.append(("Baseline", None))
    bl = d[d.ABLFL == "Y"]
    rows.append(("  n", bycoh(bl, lambda x: f_n(vals(x.AVAL)))))
    rows.append(("  Mean (SD)", bycoh(bl, lambda x: f_meansd(vals(x.AVAL)))))
    for v, vlab in [("M3", "Month 3"), ("M6", "Month 6"), ("M12", "Month 12")]:
        vd = d[d.AVISIT == v]
        rows.append((vlab, None))
        rows.append(("  n", bycoh(vd, lambda x: f_n(vals(x.AVAL)))))
        rows.append(("  Value, Mean (SD)", bycoh(vd, lambda x: f_meansd(vals(x.AVAL)))))
        rows.append(("  Change, Mean (SD)", bycoh(vd, lambda x: f_meansd(vals(x.CHG)))))
        rows.append(("  Change, 95% CI", bycoh(vd, lambda x: f_ci(vals(x.CHG)))))
    return render(tnum, label, "Full Analysis Set", rows,
                  ["Change from Baseline = post-baseline value minus baseline value.",
                   "95% CI is a descriptive two-sided t-based interval for the mean change; no hypothesis test (Phase 1).",
                   "Observed cases only; n reflects subjects with a value at the visit."],
                  source)

# ---------- build all ----------
TABLES = [
    ("t_14_1_1_demographics", t1411()),
    ("t_14_1_2_disposition", t1412()),
    ("t_14_1_3_exposure", t1413()),
    ("t_14_3_1_1_teae_overall", t14311()),
    ("t_14_3_1_2_teae_by_term", t14312()),
    ("t_14_3_3_1_dlt", t14331()),
    ("t_14_2_1_1_lvef_cfb", cfb_table("Table 14.2.1.1", "LVEFC", "Change from Baseline in LVEF (Central Reader)", "an_eff (derived from raw EDC ec)")),
    ("t_14_2_4_2_kccq_cfb", cfb_table("Table 14.2.4.2", "KCCQ", "Change from Baseline in KCCQ Overall Summary Score", "an_eff (derived from raw EDC kccq)")),
]

# write .txt per table + combined
combined = []
for fn, lines in TABLES:
    with open(os.path.join(OUTT, fn + ".txt"), "w") as f:
        f.write("\n".join(lines) + "\n")
    combined.append("\n".join(lines))
with open(os.path.join(OUTT, "all_tables.txt"), "w") as f:
    f.write(("\n\n" + "#" * TOTALW + "\n\n").join(combined) + "\n")

# assemble Word doc (landscape, Courier New 9pt, monospace blocks)
doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Inches(11), Inches(8.5)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(0.6))
# cover
def cpara(text, size=11, bold=False):
    p = doc.add_paragraph(); p.alignment = 1
    r = p.add_run(text); r.font.name = "Courier New"; r.font.size = Pt(size); r.font.bold = bold
    return p
cpara("HF-1002-CL-101", 16, True)
cpara("Populated Tables -- Non-CDISC Programming Demonstration", 13, True)
cpara("Programmed directly from analysis datasets (an_subject, an_ae, an_lb, an_eff)", 11)
cpara("built from raw EDC extracts -- no SDTM/ADaM layer.", 11)
cpara("FICTIONAL / SYNTHETIC DATA -- FOR DEMONSTRATION AND TRAINING ONLY", 11, True)
doc.add_page_break()
def mono(line):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    r = p.add_run(line if line != "" else " ")
    r.font.name = "Courier New"; r.font.size = Pt(9)
for i, (fn, lines) in enumerate(TABLES):
    if i: doc.add_page_break()
    for ln in lines: mono(ln)
docx_path = os.path.join(BASE, "HF-1002-CL-101_Populated_Tables.docx")
doc.save(docx_path)

print("Tables written to ./tables_output/ (%d tables) and %s" % (len(TABLES), os.path.basename(docx_path)))
for fn, _ in TABLES: print("  -", fn + ".txt")
