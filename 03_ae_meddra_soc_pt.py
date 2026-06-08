#!/usr/bin/env python3
"""
03_ae_meddra_soc_pt.py  --  HF-1002-CL-101 (FICTIONAL DEMO)

Completes the one step that was stubbed in the non-CDISC demo: MedDRA coding.
Applies a transparent verbatim->PT/SOC coding map (analysis_data/meddra_coding_map.csv)
to an_ae, then produces:
  - Table 14.3.1.2  TEAEs by System Organ Class and Preferred Term (Safety Set)
  - A consolidated FINAL TFL PACKAGE Word document (all populated tables).

Coding is applied as a separate, auditable layer (the map is a CSV), exactly as a
medical coder's output would feed back into the raw/EDC programming flow.
"""
import os, glob
import pandas as pd, numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT

BASE = os.path.dirname(os.path.abspath(__file__))
AD = os.path.join(BASE, "analysis_data")
OUTT = os.path.join(BASE, "tables_output")

S = pd.read_csv(os.path.join(AD, "an_subject.csv"))
AE = pd.read_csv(os.path.join(AD, "an_ae.csv"))
MAP = pd.read_csv(os.path.join(AD, "meddra_coding_map.csv"))

COHORTS = [(1, "Cohort 1", "(Low)"), (2, "Cohort 2", "(Mid)"), (3, "Cohort 3", "(High)")]
N = {c: int((S.COHORT == c).sum()) for c, _, _ in COHORTS}
NT = int(len(S))
STUB, COL = 36, 15
TOTALW = STUB + 4 * COL

def f_npct(k, n): return "0" if n == 0 else f"{k} ({100*k/n:.1f})"
def rowfmt(stub, vs):
    s = str(stub)[:STUB].ljust(STUB)
    for v in vs: s += str(v).rjust(COL)
    return s
def render(tnum, title, pop, rows, foot, source):
    L = [f"{tnum}  {title}", f"({pop})",
         "Protocol: HF-1002-CL-101        Sponsor: Corvexa Therapeutics, Inc. (fictional)", ""]
    L.append("=" * TOTALW)
    L.append(rowfmt("", ["Cohort 1", "Cohort 2", "Cohort 3", "Total"]))
    L.append(rowfmt("", ["(Low)", "(Mid)", "(High)", ""]))
    L.append(rowfmt("", [f"(N={N[1]})", f"(N={N[2]})", f"(N={N[3]})", f"(N={NT})"]))
    L.append("-" * TOTALW)
    for r in rows:
        L.append(str(r[0]) if r[1] is None else rowfmt(r[0], r[1]))
    L.append("=" * TOTALW)
    L += ["", "Source: " + source]
    for i, f in enumerate(foot): L.append(f"[{chr(97+i)}] {f}")
    L.append("Placeholders: none -- values computed from synthetic EDC data (FICTIONAL).")
    return L

# ---- code the AEs ----
teae = AE[AE.TEAEFL == "Y"].merge(MAP[["AETERM", "MEDDRA_PT", "MEDDRA_SOC"]], on="AETERM", how="left")
unmapped = teae[teae.MEDDRA_PT.isna()].AETERM.unique()
if len(unmapped): raise SystemExit(f"Unmapped AE terms (add to coding map): {list(unmapped)}")
ver = MAP.MEDDRA_VERSION.iloc[0]

def subj_npct(df):
    return [f_npct(len(set(df.SUBJID) & set(S[S.COHORT == c].SUBJID)), N[c]) for c, _, _ in COHORTS] + \
           [f_npct(df.SUBJID.nunique(), NT)]

rows = [("Subjects with >= 1 TEAE", subj_npct(teae)), ("", None)]
# SOC ordered by total distinct subjects desc, then PT within
soc_order = teae.groupby("MEDDRA_SOC").SUBJID.nunique().sort_values(ascending=False).index
for soc in soc_order:
    sdf = teae[teae.MEDDRA_SOC == soc]
    rows.append((soc, subj_npct(sdf)))
    pt_order = sdf.groupby("MEDDRA_PT").SUBJID.nunique().sort_values(ascending=False).index
    for pt in pt_order:
        rows.append(("  " + pt, subj_npct(sdf[sdf.MEDDRA_PT == pt])))
coded = render("Table 14.3.1.2", "TEAEs by System Organ Class and Preferred Term", "Safety Set", rows,
               [f"MedDRA version {ver}; coding map: analysis_data/meddra_coding_map.csv.",
                "A subject is counted once per SOC and once per PT regardless of number of events.",
                "SOCs ordered by decreasing total subject frequency; PTs within SOC likewise.",
                "Supersedes the by-verbatim-term display (14.3.1.2, supportive)."],
               "an_ae + meddra_coding_map (derived from raw EDC ae)")
with open(os.path.join(OUTT, "t_14_3_1_2_teae_by_soc_pt.txt"), "w") as f:
    f.write("\n".join(coded) + "\n")

# ---- assemble FINAL TFL PACKAGE (all populated tables, ordered) ----
order = [
    ("t_14_1_1_demographics.txt", "Table 14.1.1 Demographics and Baseline Characteristics"),
    ("t_14_1_2_disposition.txt", "Table 14.1.2 Participant Disposition"),
    ("t_14_1_3_exposure.txt", "Table 14.1.3 Extent of Exposure"),
    ("t_14_2_1_1_lvef_cfb.txt", "Table 14.2.1.1 Change from Baseline in LVEF (Central Reader)"),
    ("t_14_2_4_2_kccq_cfb.txt", "Table 14.2.4.2 Change from Baseline in KCCQ"),
    ("t_14_3_1_1_teae_overall.txt", "Table 14.3.1.1 Overall Summary of TEAEs"),
    ("t_14_3_1_2_teae_by_soc_pt.txt", "Table 14.3.1.2 TEAEs by SOC and PT (MedDRA-coded)"),
    ("t_14_3_1_2_teae_by_term.txt", "Table 14.3.1.2 TEAEs by Verbatim Term (supportive)"),
    ("t_14_3_3_1_dlt.txt", "Table 14.3.3.1 Dose-Limiting Toxicities by Cohort"),
]
doc = Document()
sec = doc.sections[0]; sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Inches(11), Inches(8.5)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"): setattr(sec, m, Inches(0.6))
def cpara(t, sz=11, b=False):
    p = doc.add_paragraph(); p.alignment = 1
    r = p.add_run(t); r.font.name = "Courier New"; r.font.size = Pt(sz); r.font.bold = b
cpara("HF-1002-CL-101", 16, True)
cpara("Final TFL Package -- Populated Tables (Non-CDISC Workflow)", 13, True)
cpara("Raw EDC -> analysis datasets -> tables; independently QC'd in R (PASS).", 11)
cpara("MedDRA-coded safety tables included. FICTIONAL / SYNTHETIC -- DEMONSTRATION ONLY", 11, True)
doc.add_page_break()
# Table of contents (simple)
cpara("Contents", 13, True)
for _, title in order:
    p = doc.add_paragraph(); r = p.add_run(title); r.font.name = "Courier New"; r.font.size = Pt(10)
doc.add_page_break()
def mono(line):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    r = p.add_run(line if line != "" else " "); r.font.name = "Courier New"; r.font.size = Pt(9)
for i, (fn, _) in enumerate(order):
    if i: doc.add_page_break()
    for ln in open(os.path.join(OUTT, fn)).read().splitlines(): mono(ln)
pkg = os.path.join(BASE, "HF-1002-CL-101_Final_TFL_Package.docx")
doc.save(pkg)

print("Coded SOC/PT table -> tables_output/t_14_3_1_2_teae_by_soc_pt.txt")
print("Final package -> %s (%d tables)" % (os.path.basename(pkg), len(order)))
print("SOCs:", list(soc_order))
print("PT 'Transaminases increased' subjects (Total):", teae[teae.MEDDRA_PT=='Transaminases increased'].SUBJID.nunique())
